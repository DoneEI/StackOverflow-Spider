import ssl
import requests
from bs4 import BeautifulSoup
from docx import Document
import re
import xlrd

ssl._create_default_https_context = ssl._create_unverified_context

# Request Header
headers = {
    'Connection': 'keep-alive',
    'Accept-Encoding': 'gzip, deflate, br',
    'Cache-Control': 'no-cache',
    'Upgrade-Insecure-Requests': '1',
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/70.0.3538.110 Safari/537.36',
}

ILLEGAL_CHARACTERS_RE = re.compile(r'[\000-\010]|[\013-\014]|[\016-\037]')

# Configuration
config = {}


def request(request_url):
    try:
        r = requests.get(url=request_url, headers=headers, timeout=20)

        bs = BeautifulSoup(r.text, "html.parser")

        table = bs.select('.container')[0]

        return table
    except Exception as e:
        print('访问url: ' + request_url + ' error: ' + str(e))


def wDoc(document, order, request_url, title, question_content, question_comment, answers, answer_cor_comments):
    # 写url
    p_r = document.add_heading(str(order) + '.' + request_url, level=1)

    # 写标题
    p_t = document.add_paragraph('')
    p_t.add_run('T:').bold = True
    p_t.add_run(title)

    # 写问题内容
    p_q = document.add_paragraph('')
    p_q.add_run('Q:').bold = True
    p_q.add_run(question_content)

    # 写问题评论
    if question_comment is not None and len(question_comment) != 0:
        for j in range(len(question_comment)):
            c = document.add_paragraph('')
            c.add_run('C' + str(j + 1) + ':').bold = True
            c.add_run(question_comment[j])

    # 写回答及其评论
    ac = document.add_paragraph(str(len(answers)))
    ac.add_run(' Answer').bold = True

    for i in range(len(answers)):
        # 写回答i内容
        p = document.add_paragraph('')
        p.add_run('A' + str(i + 1) + ':').bold = True
        p.add_run(answers[i])

        # 写回答i评论
        if answer_cor_comments[i] is not None and len(answer_cor_comments[i]) != 0:
            comments = answer_cor_comments[i]
            for j in range(len(comments)):
                c = document.add_paragraph('')
                c.add_run('C' + str(j + 1) + ':').bold = True
                c.add_run(comments[j])

    document.add_page_break()


def warn_tag(tag):
    warn_tags = config['UNPROCESSABLE_TAGS_WARN_MSG']
    warn_bool_map = {}

    for t in warn_tags:
        warn_bool_map[t] = False

    warnMsg = 'WARN: THIS PARAGRAPH CONTAINS TAG: '

    for t in tag.contents:
        if type(t).__name__ == 'Tag':
            tagName = t.name
            if tagName in warn_tags:
                if not warn_bool_map[tagName]:
                    warn_bool_map[tagName] = True
                    warnMsg += '[' + str(tagName).upper() + '] '

    for v in warn_bool_map.values():
        if v:
            return True, warnMsg

    return False, ''


def get_post_content(tags):
    res = ''

    if tags.contents is None:
        return ''

    for tag in tags.contents:
        if type(tag).__name__ == 'Tag':
            if tag.name == 'p' or tag.name == 'pre':

                text = str(tag.get_text()).replace('\n', '')
                text = re.sub(' +', ' ', text)

                warn, warn_msg = warn_tag(tag)

                if warn:
                    text += '\r\n' + warn_msg
                res += text + '\r\n'

            elif tag.name == 'a':
                res += tag.get_text() + '(' + tag.get('href') + ')'
            elif tag.name == 'li':
                res += '   ● ' + tag.get_text() + '\r\n'
            elif tag.contents is not None:
                res += get_post_content(tag)

    return res


def get_post_comment(tag):
    comments = []

    for c in tag.select('.comment-copy'):
        comments.append(c.get_text())

    return comments


def save_page_info(document, order, question_url):
    table = request(question_url)

    # question标题
    title_related = table.select('.question-hyperlink')[0]
    title = title_related.get_text()

    # 获取question内容
    question_related = table.select('#question')[0]
    post_text = question_related.select('.s-prose,.js-post-body')[0]

    # 获取question 内容及评论
    question_content = get_post_content(post_text)
    question_comment = get_post_comment(question_related)

    # 获取answers内容
    answer_related = table.select('#answers')[0]
    answers = []
    answers_cor_comment = {}
    a_count = 0

    for a in answer_related.contents:
        if type(a).__name__ == 'Tag':
            if a.name == 'div' and (a.get('class') is not None and 'answer' in a.get('class')):
                answer_content = get_post_content(a.select('.s-prose,.js-post-body')[0])
                answer_comment = get_post_comment(a)

                answers.append(answer_content)
                answers_cor_comment[a_count] = answer_comment

                a_count += 1

    wDoc(document, order, question_url, title, question_content, question_comment, answers, answers_cor_comment)
    print('Url order:' + str(order) + ' OK!')


def spy_run():
    # 设置配置
    baseConfig()

    # 检查配置
    checkConfig()

    # 获取存放url的excel文档
    data = xlrd.open_workbook(config['COLLECTED_URL_EXCEL_FILE_PATH'])

    for work in range(config['NUMBER_OF_SPY_WORK']):
        # Word Doc
        document = Document()

        print("Current running work: " + str(work + 1))

        sheet = data.sheet_by_index(config['COLLECTED_URL_EXCEL_FILE_SHEET'][work])

        start_row = int(config['COLLECTED_URL_EXCEL_FILE_START_ROW'][work])
        end_row = sheet.nrows

        if 0 < config['COLLECTED_URL_EXCEL_FILE_END_ROW'][work] < sheet.nrows:
            end_row = int(config['COLLECTED_URL_EXCEL_FILE_END_ROW'][work])

        for row in range(start_row, end_row):
            try:
                url = sheet.cell(row, config['COLLECTED_URL_EXCEL_FILE_COL'][work]).value
                save_page_info(document, row + 1, url)
            except Exception:
                print('Url order:' + str(row + 1) + ' ERROR!')
                continue

        # 保存爬取的word文档 文件名可修改
        doc_name = str(sheet.name) + '.docx'
        try:
            doc_name = str(config['COLLECTED_DATA_WORD_FILE_NAME'][work])
            if not doc_name.endswith('.docx'):
                doc_name += '.docx'
        except IndexError:
            pass

        # 默认保存在Extracted Documents文件夹下
        document.save('Extracted Documents/' + doc_name)

        print(doc_name + ' Saved OK!')
        print("work: " + str(work + 1) + " Completed!")
        print("------------------------------------------------- \n")


def baseConfig():
    # url文档路径, 只支持excel文件
    config['COLLECTED_URL_EXCEL_FILE_PATH'] = 'Url Demo.xlsx'

    # 爬取一个excel文件的文档数量,即需要爬取的该excel文件工作表个数
    config['NUMBER_OF_SPY_WORK'] = 1

    # 需要爬取的工作表索引
    config['COLLECTED_URL_EXCEL_FILE_SHEET'] = [0]

    # url所在文档工作表的列, 默认全为0
    config['COLLECTED_URL_EXCEL_FILE_COL'] = [0]

    # 爬取url起始行, 默认全为0
    config['COLLECTED_URL_EXCEL_FILE_START_ROW'] = [0]

    # 爬取url结束行, 默认全为-1（表示文档末尾)
    config['COLLECTED_URL_EXCEL_FILE_END_ROW'] = [-1]

    # 保存爬取文件的文件名(没有则按工作表名命名)
    config['COLLECTED_DATA_WORD_FILE_NAME'] = ['Demo.docx']

    # 对于无法处理且需要警告的标签
    config['UNPROCESSABLE_TAGS_WARN_MSG'] = ['code', 'img']


def checkConfig():
    expect_len = config['NUMBER_OF_SPY_WORK']

    # 检查索引个数, 若不符直接抛出异常
    if expect_len != len(config['COLLECTED_URL_EXCEL_FILE_SHEET']):
        raise Exception('爬取文件总数与索引配置长度不符')

    # 检索列,起始行,结束行,若不符合长度，填充默认值

    tmp_len = len(config['COLLECTED_URL_EXCEL_FILE_COL'])
    if expect_len != tmp_len:
        for i in range(expect_len - tmp_len):
            config['COLLECTED_URL_EXCEL_FILE_COL'].append(0)

    tmp_len = len(config['COLLECTED_URL_EXCEL_FILE_START_ROW'])
    if expect_len != tmp_len:
        for i in range(expect_len - tmp_len):
            config['COLLECTED_URL_EXCEL_FILE_START_ROW'].append(0)

    tmp_len = len(config['COLLECTED_URL_EXCEL_FILE_END_ROW'])
    if expect_len != tmp_len:
        for i in range(expect_len - tmp_len):
            config['COLLECTED_URL_EXCEL_FILE_END_ROW'].append(-1)


if __name__ == '__main__':
    spy_run()
